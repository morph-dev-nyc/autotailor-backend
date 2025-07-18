from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from openai import OpenAI
from docx import Document as DocxDocument
import io
import os
import json
import re
from dotenv import load_dotenv
from tempfile import NamedTemporaryFile

load_dotenv()

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def generate_structured_resume(resume_text: str, job_description: str) -> dict:
    prompt = f"""
You are an expert resume editor and ATS optimization specialist.

Your task is to tailor the ORIGINAL RESUME below to better match the JOB DESCRIPTION. You must preserve all factual content — do not invent experience — but you should actively rephrase and reorganize the resume to highlight relevant experience, skills, and terminology from the job posting.

Focus on these goals:
- Emphasize any relevant experience or tools mentioned in the job posting (only if they already exist in the resume)
- Include ATS-friendly keywords from the job description throughout the resume (where factually applicable)
- Improve clarity, action orientation, and alignment with ATS keyword matching
- Reorganize content for readability and impact
- If the resume does not include an exact match for a job duty, highlight transferable skills and rework related experience to better align with the job responsibilities

⚠️ NEVER fabricate tools, platforms, companies, or duties not already present in the resume.
✅ If the job description includes responsibilities not mentioned in the resume, generalize or rephrase existing experience to reflect transferable skills (e.g., if the resume says “configured internal software” and the job says “configure Activu software,” emphasize relevance without adding Activu).

JOB DESCRIPTION:
{job_description}

ORIGINAL RESUME:
{resume_text}

Output valid JSON using this structure:
- "contact": string
- "summary": short paragraph (include relevant job title and years of experience)
- "skills": newline-separated bullets
- "experience": list of objects ("title", "company", "date", "bullets")
- "education": string
- "certifications": string (optional)
- "highlights": list of 3 to 6 bullet points summarizing the most impactful tailoring changes made

Return only valid JSON. No explanations or headers.
"""

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a resume editing assistant that optimizes content for job alignment and keyword relevance."},
            {"role": "user", "content": prompt}
        ]
    )

    content = response.choices[0].message.content.strip()
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        return {"error": "Failed to parse GPT response as JSON."}

def create_formatted_docx(structured: dict) -> bytes:
    doc = DocxDocument()

    if contact := structured.get("contact"):
        doc.add_paragraph(contact)

    if summary := structured.get("summary"):
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(summary)

    if skills := structured.get("skills"):
        doc.add_heading("Skills", level=2)
        for skill in skills.split("\n"):
            if skill.strip():
                doc.add_paragraph(skill.strip(), style='List Bullet')

    if experience := structured.get("experience"):
        doc.add_heading("Experience", level=2)
        for role in experience:
            doc.add_paragraph(f"{role.get('title')} – {role.get('company')} ({role.get('date')})")
            for bullet in role.get("bullets", [])[:5]:
                if bullet.strip():
                    doc.add_paragraph(bullet.strip(), style='List Bullet')

    if education := structured.get("education"):
        doc.add_heading("Education", level=2)
        doc.add_paragraph(education)

    if certs := structured.get("certifications"):
        if certs.strip():
            doc.add_heading("Certifications", level=2)
            doc.add_paragraph(certs)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def sanitize_header(value: str) -> str:
    return ''.join(c for c in value if 32 <= ord(c) < 127)

def extract_name_from_contact(contact: str) -> str:
    lines = [line.strip() for line in contact.splitlines() if line.strip()]
    if lines:
        possible_name = lines[0].split("|")[0].split(",")[0].strip()
        if re.match(r"^[A-Z][a-z]+ [A-Z][a-z]+$", possible_name):
            return possible_name
    match = re.search(r"\b([A-Z][a-z]+) ([A-Z][a-z]+)\b", contact)
    if match:
        return f"{match.group(1)} {match.group(2)}"
    return ""

def extract_name_from_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines[:10]:
        clean_line = re.sub(r"[^\w\s\-]", "", line)
        match = re.match(r"^([A-Z][a-zA-Z]+)[\s\-]+([A-Z][a-zA-Z]+)$", clean_line)
        if match:
            return f"{match.group(1)} {match.group(2)}"
    return "Tailored"

@app.post("/tailor-file")
async def tailor_file(
    file: UploadFile = File(...),
    job_description: str = Form(...),
    linkedin_company: str = Form(""),
    linkedin_title: str = Form("")
):
    content = await file.read()

    if file.filename.endswith(".docx"):
        with NamedTemporaryFile(delete=True, suffix=".docx") as tmp:
            tmp.write(content)
            tmp.flush()
            doc = DocxDocument(tmp.name)
            resume_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    else:
        resume_text = content.decode("utf-8", errors="ignore")

    structured_resume = generate_structured_resume(resume_text, job_description)
    if "error" in structured_resume:
        return {"error": structured_resume["error"]}

    contact = structured_resume.get("contact", "")
    full_name = extract_name_from_contact(contact)
    if not full_name:
        full_name = extract_name_from_text(resume_text)

    safe_name = sanitize_header(full_name.replace(" ", "_")) if full_name else "Tailored"
    safe_title = sanitize_header(linkedin_title.replace(" ", "_")) if linkedin_title else ""
    safe_company = sanitize_header(linkedin_company.replace(" ", "_")) if linkedin_company else ""

    filename_parts = [safe_name]
    if safe_title:
        filename_parts.append(safe_title)
    if safe_company:
        filename_parts.append(safe_company)

    filename = "_".join(filename_parts) + ".docx"
    highlights = structured_resume.get("highlights", [])

    buffer = create_formatted_docx(structured_resume)

    return JSONResponse(
        content={
            "filename": filename,
            "highlights": highlights,
            "file": buffer.getvalue().hex()
        }
    )

@app.post("/suggest-titles")
async def suggest_titles(file: UploadFile = File(...)):
    content = await file.read()
    text = content.decode("utf-8", errors="ignore")

    # ✅ Truncate resume content to prevent token overflow
    truncated_text = text[:3000]

    prompt = f"""
You are a career coach and job market expert.
Based on the resume below, suggest 5 to 8 realistic U.S. job titles that the candidate is qualified for.
These should reflect the candidate's current skill set, experience, and certifications, and can come from any industry or domain (not just IT).
Return them as a JSON list of strings. Do NOT include explanations or categories.

Resume:
{truncated_text}
"""

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You suggest alternate job titles based on resume text."},
            {"role": "user", "content": prompt}
        ]
    )

    content = response.choices[0].message.content.strip()
    try:
        titles = json.loads(content)
        return {"titles": titles}
    except json.JSONDecodeError:
        return {"titles": []}
